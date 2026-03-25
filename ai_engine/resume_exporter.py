import os
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT, TA_CENTER


class ResumeExporter:
    """Exports a candidate profile + accepted rewritten bullets to DOCX or PDF."""

    def build_resume_text_sections(self, profile, accepted_bullets):
        """
        Returns a dict with structured sections for the resume.
        accepted_bullets: list of {"original": ..., "rewritten": ...}
        """
        name = profile.get("name", "Your Name")
        contact = profile.get("contact", {})
        contact_line = " | ".join(filter(None, [
            contact.get("email", ""),
            contact.get("phone", ""),
            contact.get("linkedin", ""),
            contact.get("github", ""),
        ]))

        return {
            "name": name,
            "contact": contact_line,
            "summary": profile.get("summary", ""),
            "experience": [b["rewritten"] for b in accepted_bullets],
            "original_experience": profile.get("experience", []),
            "education": profile.get("education", []),
            "skills": profile.get("skills", []),
            "projects": profile.get("projects", []),
            "certifications": profile.get("certifications", []),
        }

    def export_to_docx(self, profile, accepted_bullets, output_path):
        sections = self.build_resume_text_sections(profile, accepted_bullets)

        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Name
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(sections["name"])
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x4E)  # Navy

        # Contact
        if sections["contact"]:
            contact_para = doc.add_paragraph(sections["contact"])
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in contact_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_paragraph()  # spacer

        # Summary
        if sections["summary"]:
            self._add_section_heading(doc, "PROFESSIONAL SUMMARY")
            doc.add_paragraph(sections["summary"])

        # Experience
        if sections["experience"]:
            self._add_section_heading(doc, "EXPERIENCE")
            for bullet in sections["experience"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(bullet).font.size = Pt(10)

        # Education
        if sections["education"]:
            self._add_section_heading(doc, "EDUCATION")
            for line in sections["education"]:
                doc.add_paragraph(line)

        # Skills
        if sections["skills"]:
            self._add_section_heading(doc, "SKILLS")
            skills_text = " • ".join(sections["skills"])
            doc.add_paragraph(skills_text)

        # Projects
        if sections["projects"]:
            self._add_section_heading(doc, "PROJECTS")
            for proj in sections["projects"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(proj).font.size = Pt(10)

        # Certifications
        if sections["certifications"]:
            self._add_section_heading(doc, "CERTIFICATIONS")
            for cert in sections["certifications"]:
                doc.add_paragraph(cert)

        doc.save(output_path)
        return output_path

    def export_to_pdf(self, profile, accepted_bullets, output_path):
        sections = self.build_resume_text_sections(profile, accepted_bullets)

        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        styles = getSampleStyleSheet()
        story = []

        # Custom styles
        name_style = ParagraphStyle("name", fontSize=20, fontName="Helvetica-Bold",
                                    textColor=(0.12, 0.16, 0.31), alignment=TA_CENTER, spaceAfter=4)
        contact_style = ParagraphStyle("contact", fontSize=9, textColor=(0.39, 0.46, 0.55),
                                       alignment=TA_CENTER, spaceAfter=12)
        section_style = ParagraphStyle("section", fontSize=10, fontName="Helvetica-Bold",
                                       textColor=(0.08, 0.51, 0.82), spaceBefore=12, spaceAfter=4,
                                       borderPad=2)
        body_style = ParagraphStyle("body", fontSize=10, leading=14, spaceAfter=4)
        bullet_style = ParagraphStyle("bullet", fontSize=10, leading=14, leftIndent=12,
                                      bulletIndent=0, spaceAfter=2)

        story.append(Paragraph(sections["name"], name_style))
        if sections["contact"]:
            story.append(Paragraph(sections["contact"], contact_style))

        def section_heading(title):
            story.append(Paragraph(f"<u>{title}</u>", section_style))

        if sections["summary"]:
            section_heading("PROFESSIONAL SUMMARY")
            story.append(Paragraph(sections["summary"], body_style))

        if sections["experience"]:
            section_heading("EXPERIENCE")
            for bullet in sections["experience"]:
                story.append(Paragraph(f"• {bullet}", bullet_style))
            story.append(Spacer(1, 6))

        if sections["education"]:
            section_heading("EDUCATION")
            for line in sections["education"]:
                story.append(Paragraph(line, body_style))

        if sections["skills"]:
            section_heading("SKILLS")
            story.append(Paragraph(" • ".join(sections["skills"]), body_style))

        if sections["projects"]:
            section_heading("PROJECTS")
            for proj in sections["projects"]:
                story.append(Paragraph(f"• {proj}", bullet_style))

        if sections["certifications"]:
            section_heading("CERTIFICATIONS")
            for cert in sections["certifications"]:
                story.append(Paragraph(cert, body_style))

        doc.build(story)
        return output_path

    def _add_section_heading(self, doc, title):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x14, 0x82, 0xD1)  # Blue

        # Add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1482D1")
        pBdr.append(bottom)
        pPr.append(pBdr)
